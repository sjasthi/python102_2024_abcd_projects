# !pip install python-docx
from docx import Document
import json
import requests

# every abcd id that has the tag "sheroes"
abcd_ids = [26, 27, 28, 29, 30, 31, 32, 33, 39, 50, 52, 53, 101, 102, 110, 111, 112, 114, 115, 116, 117, 119, 151,
            171, 183, 188, 196, 201, 206, 235, 265, 275, 276, 306, 313, 314, 316, 317, 318, 319, 320, 321, 322, 323,
            324, 325, 326, 327, 328, 329, 330, 356, 401, 405, 406, 407, 409, 410, 411, 412, 413, 414, 415, 418, 419,
            420, 422, 423, 424, 425, 426, 427, 428, 429, 430, 431, 432, 433, 434, 435, 436, 437, 438, 439, 440, 441,
            442, 443, 444, 445, 462, 463, 464, 465, 466, 467, 468, 469, 470, 471, 472, 473, 474, 475, 476, 477, 478,
            483, 484, 491, 492, 493, 502, 506, 518, 520, 523, 524, 525, 534, 538, 542, 544, 547, 548, 549, 560, 568,
            574, 575, 577, 578, 581, 582, 583, 586, 588, 601, 605, 611, 618, 620, 626, 631, 632, 654, 655, 658, 659,
            660, 662, 663, 664, 665, 666, 667, 668, 669, 670, 671, 672, 673, 674, 676, 677, 678, 679, 680, 681, 682,
            683, 684, 685, 686, 687, 688, 689, 690, 691, 692, 693, 694, 695, 696, 697, 698, 699, 700, 701, 702, 703,
            704, 706, 709, 710, 711, 712, 713, 714, 715, 716, 717, 718, 719, 720, 721, 722, 723, 724, 725, 726, 728,
            729, 730, 731, 732, 733, 734, 735, 736, 737, 738, 739, 740, 741, 742, 743, 744, 745, 746, 747, 748, 749,
            750, 751, 752, 753, 754, 755, 756, 757, 758, 759, 760, 761, 762, 763, 764, 765, 766, 767, 768, 769, 770,
            771, 772]
# test ids for the presentation
test_abcd_ids = [26, 27, 28, 29, 30, 31, 32, 33, 39, 50]

def review_document(sheroes: list):
    # initialize the document as a method
    abcd_document = Document()
    # creating each dictionary and adding it to an empty list
    for x in sheroes:
        # load all of the necessary data
        response = requests.get(f'https://abcd2.projectabcd.com/api/getinfo.php?id={x}', headers={"User-Agent": "XY"})
        json_response = dict(json.loads(response.text))
        # response_code basically tells us whether any data exists for a certain id or not
        if json_response['response_code'] == 400:
            continue
        # add the name and abcd id as a document title
        abcd_document.add_heading(f'{json_response['data']['name']} ABCD ID: {json_response['data']['id']}', level=0)
        # add the description and did_you_know
        abcd_document.add_paragraph(f'Main Description: \n{json_response['data']['description']}')
        abcd_document.add_paragraph(f'Did you know: \n{json_response['data']['did_you_know']}')
        # technical details are accounted for here
        abcd_document.add_heading('Technical Stuff: \n', level=1)
        abcd_document.add_paragraph(
            f'Category: {json_response['data']['category']}\nType: {json_response['data']['type']}\nState Abbreviation: {json_response['data']['state_name']}\nKey Words: {json_response['data']['key_words']}\nCurrent Status: {json_response['data']['status']}\nNotes: {json_response['data']['notes']}\nBook: {json_response['data']['book']}\nTag Line: {json_response['data']['tag_line']}',
            style='List Bullet')
        # newline doesn't work for doc or docx, so we page break instead
        abcd_document.add_page_break()
    # saves all of these methods to a new file
    abcd_document.save(f'Project_ABCD_ID:{sheroes[0]}, ..., {sheroes[-1]}.docx')

# the real version would be able to do this, but it would take way too long
# review_document(abcd_ids)

# you could also go by each id, but the function would break unless it is surrounded by brackets
# review_document([753])

# we do this one instead in the interest of time
review_document(test_abcd_ids)
